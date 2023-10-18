from docx import Document
from docx.shared import Inches,Cm,Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from dataProc import DataObject
import pandas as pd
from tkinter import filedialog
import os,re, datetime

class Report():

	def __init__(self, cols_name, model_file):
		self.cols_name = cols_name
		# self.path = path
		# self.docx = Document(r'D:\试验报告\MOS可靠性试验报告模板.docx')
		# self.docx = Document(r'D:\试验报告\Diode可靠性试验报告模板.docx')
		self.docx = Document(model_file)
		self.save_path = r'D:\试验报告\可靠性试验报告.docx'

	def loadDataFromDF(self, path):
		data = DataObject(path, cols_name = self.cols_name)
		self.df = data.df
		self.project_name = list(self.df.index.get_level_values('PROJECT').unique())[0]

	def loadDataFromExcel(self, file_path):
		self.df = pd.read_excel(file_path, sheet_name = 'DATA', index_col=[0,1,2], header = [0,1])
		self.project_name = list(self.df.index.get_level_values('PROJECT').unique())[0]

	def styleUpdate(self, paragraph):

		paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
		paragraph.runs[0].font.size = Pt(8)
		paragraph.runs[0].font.name = '宋体'
		paragraph.runs[0].font.name = 'Times New Roman'

	def fillAllTable(self):

		data_dict = {}

		for test_ID in list(self.df.index.get_level_values('TEST').unique()):
			test_symbol = re.search(r'(.*[A-Z])', test_ID).group()
			print(test_symbol)
			test_flag = '</' + test_symbol + '>'
			data_dict[test_flag] = test_ID

		for table in self.docx.tables:
			row = table.row_cells(2)
			# print(row)
			if len(row) > 0 and row[0].text in data_dict.keys():
				# print('Find a table: ', row[0].text)
				test_ID = data_dict[row[0].text]
				df = self.df.loc[(self.project_name, test_ID, slice(None)), :]
				self.__fillOneTable(table, df)

	def appendData(self, test_list):
		data_dict = test_list

		# for test_ID in test_list:
		# 	test_flag = '</' + test_ID + '>'
		# 	data_dict[test_flag] = test_ID
		# 	print(test_ID)

		for table in self.docx.tables:
			row = table.row_cells(2)
			if len(row) > 0 and row[0].text in data_dict.keys():
				# print('Find a table: ', row[0].text)
				test_ID = data_dict[row[0].text]
				df = self.df.loc[(self.project_name, test_ID, slice(None)), :]
				print(df)
				self.__fillOneTable(table, df)

	def __fillOneTable(self, table, df):

		table.style.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
		test_index = list(df.index.get_level_values('TEST').unique())[0]
		test_ID = re.search(r'(.*[A-Z])', test_index).group()
		cols_name = self.cols_name
		for i in range(len(df)-1):
			table.add_row()

		row = 2
		for index in list(df.index.get_level_values('INDEX')):

			row_cells = table.row_cells(row)
			row += 1
			
			
			row_cells[0].text = test_ID + '-' + str(index) + '#'
			self.styleUpdate(row_cells[0].paragraphs[0])

			for col_num in range(len(cols_name)):

				try:
					data1 = df.loc[(self.project_name, test_index, index), (cols_name[col_num],'PRE')]
					row_cells[col_num * 3 + 1].text = str(round(data1, 2))
				except:
					print(f'Data Error : {test_ID}-{index}# ({cols_name[col_num]}) PRE')

				try:
					data2 = df.loc[(self.project_name, test_index, index), (cols_name[col_num],'POST')]
					row_cells[col_num * 3 + 2].text = str(round(data2, 2))
				except:
					print(f'Data Error : {test_ID}-{index}# ({cols_name[col_num]}) POST')
						
				try:
					data3 = df.loc[(self.project_name, test_index, index), (cols_name[col_num],'RATE')]
					row_cells[col_num * 3 + 3].text =  str(data3)
					row_cells[col_num * 3 + 3].paragraphs[0].runs[0].font.bold = True
				except:
					print(f'Data Error : {test_ID}-{index}# ({cols_name[col_num]}) RATE')

				for i in range(col_num * 3 + 1, col_num * 3 + 4):
					if len(row_cells[i].paragraphs[0].runs) > 0:
						self.styleUpdate(row_cells[i].paragraphs[0])	

	def _insertOneImg(self, img_file, position):
		if '<<' in paragraph.text and paragraph.text[2:-2] in img_dict.keys():
			img_key = paragraph.text[2:-2]
			paragraph.text = ''
			run = paragraph.runs[0]
			# 添加图片并指定大小
			run.add_picture(img_dict[img_key], width=Cm(19))
			# print("add a pic")

	def insertImg(self, img_para):

		img_dict = {}
		if type(img_para) == tuple:
			for img in img_para:
				file = os.path.basename(img)
				img_path = os.path.dirname(img)
				img_dict[file[:-4]] = img_path + '/' + file
		else:
			for file in os.listdir(img_para):
				if file[-3:] == 'jpg':
					img_dict[file[:-4]] = img_path + '/' + file
		
		# print(img_dict)
		for paragraph in self.docx.paragraphs:
			if '<<' in paragraph.text and paragraph.text[2:-2] in img_dict.keys():
				img_key = paragraph.text[2:-2]
				paragraph.text = ''
				run = paragraph.runs[0]
				# 添加图片并指定大小
				run.add_picture(img_dict[img_key], width=Cm(19))
				# print("add a pic")

	def saveDocx(self, path):
		date = datetime.datetime.today().strftime('%Y%m%d_%H%M%S')
		self.docx.save(os.path.join(path, f'{self.project_name} -{date}.docx'))

if __name__ == "__main__":

	cols_name = [
		'BVDSS@VDSmax=1800V&IDS=200uA', 
		'VGS(TH)@IDS=10mA',
		'IDSS@VDS=1200V', 
		'IGSS@VGS=18V', 
		'RDS(ON)@IDS=40A&VG=15V']
	# cols_name = ['BVR@VRmax=1800V&IR=200uA', 'IR@VR=1200V', 'VF@IF=25A']

	# path = filedialog.askdirectory() 
	file_path = filedialog.askopenfilename()
	path = os.path.dirname(file_path)
	# model_file = r'D:\可靠性实验\试验报告\S1M040120H\S1M040120H可靠性试验报告模板.docx'
	# docx = Report(cols_name, model_file)
	# docx.loadDataFromDF(path)
	docx = Report(cols_name, r'D:\可靠性实验\试验报告\S1M040120H\S1M040120H可靠性试验报告模板.docx')
	docx.loadDataFromExcel(file_path)
	# print(docx.df)
	docx.fillAllTable()
	# docx.appendData({'</4D>':'D240'})

	# docx.insertImg(os.path.join(path, 'graph'))
	docx.insertImg(filedialog.askopenfilenames())
	docx.saveDocx(path)



